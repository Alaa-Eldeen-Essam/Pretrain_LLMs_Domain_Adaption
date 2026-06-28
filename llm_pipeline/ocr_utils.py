from __future__ import annotations

import gc
from pathlib import Path
from typing import Dict, Iterator, Optional


def pdf_page_count(path: Path, *, max_pages: Optional[int] = None) -> int:
    """Return the number of PDF pages visible to Poppler.

    Needed for terminal progress bars and page-level OCR resume. Used by the
    OCR extraction script before processing each PDF.

    Inputs: PDF path and optional max-pages cap.
    Outputs: effective page count as an integer.
    """
    from pdf2image import pdfinfo_from_path

    page_count = int(pdfinfo_from_path(str(path)).get("Pages", 0))
    if max_pages is not None:
        return min(page_count, int(max_pages))
    return page_count


def ocr_pdf_pages(
    path: Path,
    *,
    lang: str,
    dpi: int,
    max_pages: Optional[int] = None,
    start_page: int = 1,
) -> Iterator[Dict]:
    """Render and OCR PDF pages one at a time.

    Needed to avoid loading entire PDFs into memory and to resume from a
    partially completed page number. Used by the OCR extraction script.

    Inputs: PDF path, Tesseract language string, DPI, optional max-pages cap,
    and start page.
    Outputs: iterator of page OCR dictionaries containing page metadata and
    extracted text.
    """
    from pdf2image import convert_from_path
    import pytesseract

    page_count = pdf_page_count(path, max_pages=max_pages)

    # Render one page at a time. Rendering a full PDF at once can make WSL kill
    # the process under memory pressure before Python can raise an exception.
    for page_index in range(max(1, start_page), page_count + 1):
        images = convert_from_path(
            str(path),
            dpi=dpi,
            first_page=page_index,
            last_page=page_index,
            thread_count=1,
            grayscale=True,
            fmt="jpeg",
            jpegopt={"quality": 85, "progressive": False, "optimize": True},
        )
        image = images[0]
        try:
            if image.mode not in ("RGB", "L"):
                image = image.convert("RGB")
            text = pytesseract.image_to_string(image, lang=lang)
            yield {
                "page": page_index,
                "page_count": page_count,
                "ocr_engine": "tesseract",
                "ocr_lang": lang,
                "dpi": dpi,
                "text": text or "",
                "char_count": len(text or ""),
            }
        finally:
            image.close()
            del images
            gc.collect()


def ocr_image(path: Path, *, lang: str) -> str:
    """OCR a single image file with Tesseract.

    Needed for image sources in the corpus. Used by the OCR extraction script.

    Inputs: image path and Tesseract language string.
    Outputs: extracted text string.
    """
    from PIL import Image
    import pytesseract

    with Image.open(path) as image:
        if image.mode not in ("RGB", "L"):
            image = image.convert("RGB")
        return pytesseract.image_to_string(image, lang=lang) or ""


def read_docx(path: Path) -> str:
    """Extract text from a DOCX document.

    Needed because non-PDF text documents should not be OCRed unnecessarily.
    Used by the OCR extraction script for `.docx` files.

    Inputs: DOCX file path.
    Outputs: joined paragraph and table-cell text.
    """
    from docx import Document

    doc = Document(str(path))
    parts = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = (cell.text or "").strip()
                if text:
                    parts.append(text)
    return "\n".join(parts)
